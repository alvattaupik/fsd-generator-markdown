
def extract_use_case_name(md_content):
    m = re.search(r"\|\s*Nama Use Case\s*\|\s*([^|]+)\|", md_content, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    m = re.search(r"\*\*Tabel\s*1:\s*Deskripsi\s*(?:Detail\s*)?(?:Use\s*Case\s*)?([^*]+)\*\*", md_content, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return "Use Case Specification"

"""
Markdown to Docx Converter Engine
Version: 1.0.12
Author: Antigravity AI
Description: Converts ordered Markdown usecase files to MS Word (.docx) documents following strict FSD layout & formatting standards.
Features in v1.0.12:
- Usecase Description Table Justify Alignment: All text in the Usecase Description table (including native numbered list items for Pre-condition, Post-condition, Integrasi, Asumsi, Keterbatasan, and Aturan Bisnis/Sistem) is JUSTIFY aligned (WD_ALIGN_PARAGRAPH.JUSTIFY).
- Guaranteed Isolated AbstractNum Restart Numbering: Each section creates a dedicated abstractNum & numId with startOverride=1 in numbering.xml, forcing MS Word to restart numbering at 1 for every element section.
- Native Word Caption SEQ Fields: Captions use Word's native SEQ field XML (SEQ Tabel \\* ARABIC and SEQ Gambar \\* ARABIC) for automatic DAFTAR TABEL & DAFTAR GAMBAR indexing.
- Body Caption Renumbering: Scans template body paragraphs for highest body Tabel N and Gambar N, continuing numbering seamlessly.
- Spasi Enter Presisi: Jarak enter otomatis antara Heading & Gambar, Heading & Caption, serta Caption & Heading.
- Table Formatting: cantSplit rows, tblHeader repeating header rows, custom inner cell margins (padding), column alignment, and bold Elemen Spesifikasi.
"""

import os
import sys
import re
import argparse
import tempfile
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import requests

__version__ = '1.0.12'

def get_body_caption_max_numbers(doc):
    """Scan body paragraphs (skipping DAFTAR TABEL / DAFTAR GAMBAR) to find max Tabel N and Gambar N in the template body."""
    max_tabel = 0
    max_gambar = 0
    
    body_start_idx = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading 1') and 'Pendahuluan' in p.text:
            body_start_idx = i
            break
    
    for i in range(body_start_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        
        t_match = re.search(r'^Tabel\s+(\d+)', text, re.IGNORECASE)
        if t_match:
            num = int(t_match.group(1))
            if num > max_tabel:
                max_tabel = num
        
        g_match = re.search(r'^Gambar\s+(\d+)', text, re.IGNORECASE)
        if g_match:
            num = int(g_match.group(1))
            if num > max_gambar:
                max_gambar = num

    if max_tabel == 0:
        max_tabel = 10
    if max_gambar == 0:
        max_gambar = 2

    return max_tabel, max_gambar

def set_table_margins(table, top=120, bottom=120, left=180, right=180):
    tblPr = table._tbl.tblPr
    for m in tblPr.xpath('./w:tblCellMar'):
        tblPr.remove(m)
    tblCellMar = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(tblCellMar)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if not trPr.xpath('./w:cantSplit'):
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    if not trPr.xpath('./w:tblHeader'):
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def set_cell_shading(cell, color_hex="9BC2E6"):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.xpath('./w:shd'):
        tcPr.remove(s)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    tcPr.append(shd)

def create_fresh_abstract_and_num_id(doc):
    """
    Creates a completely fresh abstractNum AND numId with startOverride=1.
    Guarantees 100% that MS Word restarts numbering at 1 for every list section!
    """
    numbering = doc.part.numbering_part.element
    
    ab_ids = [int(ab.get(qn('w:abstractNumId'))) for ab in numbering.xpath('.//w:abstractNum')]
    num_ids = [int(num.get(qn('w:numId'))) for num in numbering.xpath('.//w:num')]
    
    new_ab_id = (max(ab_ids) if ab_ids else 100) + 1
    new_num_id = (max(num_ids) if num_ids else 100) + 1
    
    ns = nsdecls("w")
    ab_xml = f'''<w:abstractNum {ns} w:abstractNumId="{new_ab_id}">
      <w:multiLevelType w:val="hybridMultilevel"/>
      <w:lvl w:ilvl="0">
        <w:start w:val="1"/>
        <w:numFmt w:val="decimal"/>
        <w:lvlText w:val="%1."/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:ind w:left="720" w:hanging="360"/>
        </w:pPr>
      </w:lvl>
    </w:abstractNum>'''
    
    num_xml = f'''<w:num {ns} w:numId="{new_num_id}">
      <w:abstractNumId w:val="{new_ab_id}"/>
      <w:lvlOverride w:ilvl="0">
        <w:startOverride w:val="1"/>
      </w:lvlOverride>
    </w:num>'''
    
    first_num = numbering.xpath('.//w:num')
    if first_num:
        first_num[0].addprevious(parse_xml(ab_xml))
    else:
        numbering.append(parse_xml(ab_xml))
        
    numbering.append(parse_xml(num_xml))
    return new_num_id

def add_native_numPr(p, num_id):
    """Add native MS Word XML numbering (numPr) to paragraph properties."""
    pPr = p._p.get_or_add_pPr()
    for n in pPr.xpath('./w:numPr'):
        pPr.remove(n)
    numPr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="{num_id}"/></w:numPr>')
    pPr.append(numPr)

def render_plantuml(puml_text, out_img_path):
    puml_text_clean = puml_text.strip()
    if not puml_text_clean.startswith('@startuml'):
        puml_text_clean = '@startuml\n' + puml_text_clean
    if not puml_text_clean.endswith('@enduml'):
        puml_text_clean = puml_text_clean + '\n@enduml'

    try:
        resp = requests.post("https://kroki.io/plantuml/png", data=puml_text_clean.encode('utf-8'), timeout=20)
        if resp.status_code == 200 and len(resp.content) > 100:
            with open(out_img_path, 'wb') as f:
                f.write(resp.content)
            return True
    except Exception as e:
        print(f"[Warning] Kroki rendering attempt failed: {e}")

    try:
        import plantuml
        pl = plantuml.PlantUML(url='http://www.plantuml.com/plantuml/img/')
        png_data = pl.processes(puml_text_clean)
        if png_data and len(png_data) > 100:
            with open(out_img_path, 'wb') as f:
                f.write(png_data)
            return True
    except Exception as e:
        print(f"[Warning] PlantUML server rendering attempt failed: {e}")

    return False

def apply_run_formatting(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)

def apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.paragraph_format.line_spacing = line_spacing
    p.alignment = align

def add_native_caption_paragraph(doc, prefix, seq_type, number_val, title_text):
    """Add a native Word caption paragraph with SEQ field for automatic Word TOC/TOF indexing."""
    style_to_use = 'Caption' if 'Caption' in doc.styles else 'Normal'
    p = doc.add_paragraph(style=style_to_use)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    pPr = p._p.get_or_add_pPr()
    if not pPr.xpath('./w:keepNext'):
        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))

    r1 = p.add_run(prefix)
    apply_run_formatting(r1, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    
    r2 = p.add_run()
    apply_run_formatting(r2, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    r2._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    
    r3 = p.add_run()
    apply_run_formatting(r3, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    r3._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> SEQ {seq_type} \\* ARABIC </w:instrText>'))
    
    r4 = p.add_run()
    apply_run_formatting(r4, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    r4._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    
    r5 = p.add_run(str(number_val))
    apply_run_formatting(r5, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    
    r6 = p.add_run()
    apply_run_formatting(r6, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    r6._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    
    r7 = p.add_run(f" {title_text}")
    apply_run_formatting(r7, font_name="Arial", size_pt=11, color_rgb=(0,0,0))
    
    apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p

def parse_markdown_content(md_content):
    lines = md_content.splitlines()
    blocks = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            raw_title = line.lstrip('#').strip()
            clean_title = re.sub(r'^\d+(\.\d+)*\s*', '', raw_title)
            blocks.append({
                'type': 'heading',
                'level': level,
                'raw_title': raw_title,
                'title': clean_title
            })
            i += 1
            continue
        
        if line.startswith('**Tabel') or line.startswith('**Gambar'):
            caption_text = line.strip('*').strip()
            blocks.append({
                'type': 'caption',
                'text': caption_text,
                'is_table': caption_text.startswith('Tabel'),
                'is_image': caption_text.startswith('Gambar')
            })
            i += 1
            continue
        
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            headers = []
            rows = []
            for tline in table_lines:
                if '---' in tline:
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                if not headers:
                    headers = cells
                else:
                    rows.append(cells)
            
            blocks.append({
                'type': 'table',
                'headers': headers,
                'rows': rows
            })
            continue
        
        if line.startswith('```plantuml'):
            i += 1
            puml_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                puml_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            
            blocks.append({
                'type': 'plantuml',
                'code': '\n'.join(puml_lines)
            })
            continue
        
        blocks.append({
            'type': 'text',
            'text': line
        })
        i += 1

    return blocks

def parse_and_append_runs(p, text, default_bold=False):
    tokens = re.split(r'(\*\*.*?\*\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            bold_text = token[2:-2]
            run = p.add_run(bold_text)
            apply_run_formatting(run, font_name="Arial", size_pt=11, bold=True, color_rgb=(0,0,0))
        else:
            run = p.add_run(token)
            apply_run_formatting(run, font_name="Arial", size_pt=11, bold=default_bold, color_rgb=(0,0,0))

def add_cell_formatted_content(doc, cell, text, is_special_numbered_row=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, default_bold=False):
    """
    Add text to table cell.
    All text in usecase description table (including numbered list items) is JUSTIFY aligned!
    """
    p = cell.paragraphs[0]
    p.text = ""
    
    if is_special_numbered_row and ('<br>' in text or '\n' in text or re.search(r'\d+[\.\)]\s*', text)):
        # Allocate a fresh isolated abstractNum and numId with startOverride=1
        section_num_id = create_fresh_abstract_and_num_id(doc)
        
        parts = [pt.strip() for pt in re.split(r'<br\s*/?>|\n', text) if pt.strip()]
        for idx, part in enumerate(parts):
            if idx > 0:
                p = cell.add_paragraph()
            
            # Apply Native MS Word Decimal Numbering with fresh section numId
            add_native_numPr(p, num_id=section_num_id)
            
            # v1.0.12 Revision: All text in usecase description (including numbered list items) MUST be JUSTIFY aligned!
            apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            
            # Strip hardcoded number prefix ("1. ", "2. ", etc.)
            clean_part = re.sub(r'^\d+[\.\)]\s*', '', part)
            parse_and_append_runs(p, clean_part, default_bold=default_bold)
    else:
        apply_paragraph_formatting(p, line_spacing=1.5, align=align)
        parse_and_append_runs(p, text, default_bold=default_bold)

def append_blocks_to_docx(doc, blocks, temp_dir, initial_tabel_counter=10, initial_gambar_counter=2, use_case_name="Use Case Specification"):
    """Append structured markdown blocks into docx document with native Word SEQ field captions and native numPr list items."""
    special_fields = [
        'Pre-condition', 'Post-condition', 'Integrasi', 
        'Asumsi', 'Keterbatasan', 'Aturan Bisnis/Sistem', 'Aturan Bisnis', 'Aturan Sistem'
    ]
    
    tabel_counter = initial_tabel_counter
    gambar_counter = initial_gambar_counter
    img_counter = 1
    
    for b_idx, block in enumerate(blocks):
        btype = block['type']
        
        if btype == 'heading':
            raw_title = block.get('raw_title', '')
            clean_title = block['title']
            if "deskripsi use case" in raw_title.lower() or raw_title.startswith("1"):
                style_name = "Heading 4"
                heading_text = use_case_name
                p = doc.add_paragraph(heading_text, style=style_name)
                for r in p.runs:
                    apply_run_formatting(r, font_name="Arial", color_rgb=(0,0,0))
            else:
                style_name = "Heading 5"
                heading_text = clean_title
                p = doc.add_paragraph(heading_text, style=style_name)
                for r in p.runs:
                    apply_run_formatting(r, font_name="Arial", color_rgb=(0,0,0))
                # Heading 5 is always followed by enter paragraph space
                p_enter = doc.add_paragraph()
                apply_paragraph_formatting(p_enter, line_spacing=1.5)
        
        elif btype == 'caption':
            raw_text = block['text']
            if block['is_table']:
                tabel_counter += 1
                title_text = re.sub(r'^Tabel\s+\d+[:\s]*', '', raw_text, flags=re.IGNORECASE).strip()
                p = add_native_caption_paragraph(doc, "Tabel ", "Tabel", tabel_counter, title_text)
            elif block['is_image']:
                gambar_counter += 1
                title_text = re.sub(r'^Gambar\s+\d+[:\s]*', '', raw_text, flags=re.IGNORECASE).strip()
                p = add_native_caption_paragraph(doc, "Gambar ", "Gambar", gambar_counter, title_text)
            else:
                p = doc.add_paragraph(raw_text, style='Caption' if 'Caption' in doc.styles else 'Normal')
                apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)

            # If caption is followed by a heading -> add enter space
            if b_idx + 1 < len(blocks) and blocks[b_idx + 1]['type'] == 'heading':
                p_enter = doc.add_paragraph()
                apply_paragraph_formatting(p_enter, line_spacing=1.5)
        
        elif btype == 'table':
            headers = block['headers']
            rows = block['rows']
            
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            set_table_margins(table, top=120, bottom=120, left=180, right=180)
            
            hdr_row = table.rows[0]
            set_row_header(hdr_row)
            set_row_cant_split(hdr_row)
            
            for col_idx, header_text in enumerate(headers):
                cell = hdr_row.cells[col_idx]
                set_cell_shading(cell, "9BC2E6")
                p = cell.paragraphs[0]
                p.text = ""
                apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run(header_text)
                apply_run_formatting(run, font_name="Arial", size_pt=11, bold=True, color_rgb=(0,0,0))
            
            col_alignments = []
            for htext in headers:
                ht_clean = htext.strip().lower()
                if ht_clean in ['no', 'no.', 'kode', 'http status', 'response code', 'id use case']:
                    col_alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    col_alignments.append(WD_ALIGN_PARAGRAPH.JUSTIFY)
            
            is_deskripsi_usecase_table = (len(headers) == 2 and headers[0] == 'Elemen Spesifikasi')
            
            for r_idx, row_data in enumerate(rows):
                row = table.rows[r_idx + 1]
                set_row_cant_split(row)
                elemen_name = row_data[0] if len(row_data) > 0 else ""
                is_special = elemen_name in special_fields
                
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < len(row.cells):
                        cell = row.cells[c_idx]
                        align = col_alignments[c_idx] if c_idx < len(col_alignments) else WD_ALIGN_PARAGRAPH.JUSTIFY
                        is_numbered = (c_idx == 1 and is_special)
                        cell_bold = (is_deskripsi_usecase_table and c_idx == 0)
                        
                        add_cell_formatted_content(
                            doc, cell, cell_text, 
                            is_special_numbered_row=is_numbered, 
                            align=align, 
                            default_bold=cell_bold
                        )
            
            p_after = doc.add_paragraph()
            apply_paragraph_formatting(p_after, line_spacing=1.5)
        
        elif btype == 'plantuml':
            img_path = os.path.join(temp_dir, f"plantuml_diagram_{img_counter}.png")
            img_counter += 1
            
            if render_plantuml(block['code'], img_path):
                p = doc.add_paragraph()
                apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
            else:
                p = doc.add_paragraph("[PlantUML Diagram Rendering Error]")
                apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        elif btype == 'text':
            p = doc.add_paragraph()
            apply_paragraph_formatting(p, line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            parse_and_append_runs(p, block['text'])
            
    return tabel_counter, gambar_counter

def get_sorted_md_files(folder_path):
    files = [f for f in os.listdir(folder_path) if f.lower().endswith('.md')]
    def sort_key(filename):
        match = re.match(r'^(\d+)', filename)
        if match:
            return (0, int(match.group(1)), filename.lower())
        return (1, filename.lower())
    files.sort(key=sort_key)
    return [os.path.join(folder_path, f) for f in files]

def convert_md_to_docx(input_path, template_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    if os.path.isfile(input_path) and input_path.lower().endswith('.md'):
        md_files = [input_path]
    elif os.path.isdir(input_path):
        md_files = get_sorted_md_files(input_path)
    else:
        raise ValueError(f"Path invalid or does not contain markdown files: {input_path}")
    
    if not md_files:
        print(f"No markdown files found in {input_path}")
        return
    
    print(f"Found {len(md_files)} markdown file(s) to process.")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        for md_file in md_files:
            print(f"Processing: {os.path.basename(md_file)}...")
            with open(md_file, 'r', encoding='utf-8-sig') as f:
                md_content = f.read()
            
            use_case_name = extract_use_case_name(md_content)
            blocks = parse_markdown_content(md_content)
            doc = docx.Document(template_path)
            
            # Check if Detail Spesifikasi Heading 2 is already present at the end of the template
            max_tabel, max_gambar = get_body_caption_max_numbers(doc)
            
            append_blocks_to_docx(doc, blocks, temp_dir, initial_tabel_counter=max_tabel, initial_gambar_counter=max_gambar, use_case_name=use_case_name)
            
            base_name = os.path.splitext(os.path.basename(md_file))[0]
            clean_name = re.sub(r'^\d+\.\s*', '', base_name)
            out_docx_name = f"Dokumen_{clean_name}.docx"
            out_docx_path = os.path.join(output_dir, out_docx_name)
            
            os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
            doc.save(out_docx_path)
            print(f"  -> Generated: {out_docx_path}")

def main():
    parser = argparse.ArgumentParser(description=f"Convert Markdown Usecase files to formatted MS Word (.docx) document (v{__version__}).")
    parser.add_argument("input_path", help="Folder containing markdown files (e.g. 1.name.md) or path to single .md file")
    parser.add_argument("--template", required=True, help="Path to template .docx file")
    parser.add_argument("--output", required=True, help="Folder to save output .docx files")
    
    args = parser.parse_args()
    convert_md_to_docx(args.input_path, args.template, args.output)

if __name__ == '__main__':
    main()
